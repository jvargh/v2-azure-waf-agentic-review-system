"""Convert Word document to Markdown format."""
from pathlib import Path
try:
    from docx import Document
except ImportError:
    print("Installing python-docx...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

def convert_docx_to_markdown(docx_path: Path, output_path: Path = None):
    """Convert a .docx file to markdown format.
    
    Args:
        docx_path: Path to the .docx file
        output_path: Optional output path for the .md file. 
                    If None, uses same name with .md extension
    """
    if output_path is None:
        output_path = docx_path.with_suffix('.md')
    
    doc = Document(docx_path)
    markdown_lines = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        if not text:
            markdown_lines.append('')
            continue
        
        # Handle heading styles
        if paragraph.style.name.startswith('Heading'):
            level = paragraph.style.name.replace('Heading ', '')
            try:
                level_num = int(level)
                markdown_lines.append(f"{'#' * level_num} {text}")
            except ValueError:
                markdown_lines.append(f"# {text}")
        # Handle list items
        elif paragraph.style.name.startswith('List'):
            markdown_lines.append(f"- {text}")
        else:
            markdown_lines.append(text)
    
    # Handle tables
    for table in doc.tables:
        markdown_lines.append('')
        # Get headers
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        markdown_lines.append('| ' + ' | '.join(headers) + ' |')
        markdown_lines.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')
        
        # Get rows
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            markdown_lines.append('| ' + ' | '.join(cells) + ' |')
        markdown_lines.append('')
    
    # Write to file
    output_path.write_text('\n'.join(markdown_lines), encoding='utf-8')
    print(f"✓ Converted: {docx_path.name} → {output_path.name}")
    print(f"✓ Output saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    docx_file = Path("src/docs/Detailed User Flow.docx")
    
    if not docx_file.exists():
        print(f"Error: File not found: {docx_file}")
        exit(1)
    
    output_file = convert_docx_to_markdown(docx_file)
    
    # Display preview of the converted content
    print("\n" + "="*80)
    print("PREVIEW OF CONVERTED MARKDOWN")
    print("="*80)
    content = output_file.read_text(encoding='utf-8')
    print(content[:1000] + "..." if len(content) > 1000 else content)
